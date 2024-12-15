from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QLabel, QVBoxLayout, QPushButton, QTextEdit, QComboBox,
    QSpinBox, QFileDialog, QWidget, QLineEdit, QGroupBox, QHBoxLayout, QMessageBox
)
from PyQt5.QtCore import Qt
from huggingface_hub import InferenceClient
import docx
from fpdf import FPDF
from PyPDF2 import PdfReader
from docx import Document

# Hugging Face API Key
API_KEY = ""
client = InferenceClient(api_key=API_KEY)

# Functions
def generate_content(prompt, tone, length, word_limit=None, max_tokens=500):
    """Generate content using Hugging Face API."""
    if length == "Custom" and word_limit:
        formatted_prompt = f"Write in a {tone} tone and limit to {word_limit} words. {prompt}"
    else:
        formatted_prompt = f"Write in a {tone} tone and keep it {length}. {prompt}"
    messages = [{"role": "user", "content": formatted_prompt}]
    try:
        completion = client.chat.completions.create(
            model="meta-llama/Meta-Llama-3-8B-Instruct",
            messages=messages,
            max_tokens=max_tokens,
        )
        return completion.choices[0].message.content
    except Exception as e:
        return f"Error generating content: {e}"

def save_to_txt(content, filename):
    with open(filename, "w", encoding="utf-8") as file:
        file.write(content)
    return filename

def save_to_docx(content, filename):
    doc = docx.Document()
    doc.add_paragraph(content)
    doc.save(filename)
    return filename

def save_to_pdf(content, filename):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)

    # Encode content to replace unsupported characters
    try:
        encoded_content = content.encode("latin-1", "replace").decode("latin-1")
        pdf.multi_cell(0, 10, encoded_content)
    except UnicodeEncodeError:
        QMessageBox.critical(
            None, "Encoding Error", "Some characters could not be encoded into the PDF format."
        )
        return None

    pdf.output(filename)
    return filename


def read_file(file_path):
    """Read the content of a file (TXT, DOCX, or PDF)."""
    try:
        if file_path.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()
        elif file_path.endswith(".docx"):
            doc = Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        elif file_path.endswith(".pdf"):
            reader = PdfReader(file_path)
            return "\n".join([page.extract_text() for page in reader.pages])
        else:
            return "Unsupported file type."
    except Exception as e:
        return f"Error reading file: {e}"

# Main Application Class
class AIDocumentAgent(QMainWindow):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("AI Document Agent")
        self.setGeometry(100, 100, 700, 800)

        self.generated_content = ""

        self.initUI()

    def initUI(self):
        # Central Widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        # Layout
        main_layout = QVBoxLayout()
        central_widget.setLayout(main_layout)

        # Prompt Input
        prompt_group = QGroupBox("Prompt")
        prompt_layout = QVBoxLayout()
        prompt_group.setLayout(prompt_layout)

        self.prompt_input = QTextEdit()
        self.prompt_input.setPlaceholderText("Enter your prompt here...")
        prompt_layout.addWidget(self.prompt_input)

        main_layout.addWidget(prompt_group)

        # Tone and Length Selection
        options_group = QGroupBox("Options")
        options_layout = QHBoxLayout()
        options_group.setLayout(options_layout)

        # Tone Selection
        tone_layout = QVBoxLayout()
        tone_label = QLabel("Select Tone:")
        self.tone_combo = QComboBox()
        self.tone_combo.addItems(["formal", "casual", "professional", "custom"])
        self.custom_tone_input = QLineEdit()
        self.custom_tone_input.setPlaceholderText("Enter custom tone")
        self.custom_tone_input.setEnabled(False)

        self.tone_combo.currentTextChanged.connect(self.toggle_custom_tone)

        tone_layout.addWidget(tone_label)
        tone_layout.addWidget(self.tone_combo)
        tone_layout.addWidget(self.custom_tone_input)
        options_layout.addLayout(tone_layout)

        # Length Selection
        length_layout = QVBoxLayout()
        length_label = QLabel("Select Length:")
        self.length_combo = QComboBox()
        self.length_combo.addItems(["short", "medium", "long", "Custom"])
        self.word_limit_spinbox = QSpinBox()
        self.word_limit_spinbox.setRange(1, 10000)
        self.word_limit_spinbox.setValue(500)
        self.word_limit_spinbox.setEnabled(False)

        self.length_combo.currentTextChanged.connect(self.toggle_word_limit)

        length_layout.addWidget(length_label)
        length_layout.addWidget(self.length_combo)
        length_layout.addWidget(self.word_limit_spinbox)
        options_layout.addLayout(length_layout)

        main_layout.addWidget(options_group)

        # Generate and Save Buttons
        button_layout = QHBoxLayout()
        self.generate_button = QPushButton("Generate")
        self.generate_button.clicked.connect(self.generate_text)
        button_layout.addWidget(self.generate_button)

        self.save_button = QPushButton("Save")
        self.save_button.clicked.connect(self.save_text)
        self.save_button.setEnabled(False)
        button_layout.addWidget(self.save_button)

        self.upload_button = QPushButton("Upload File")
        self.upload_button.clicked.connect(self.upload_file)
        button_layout.addWidget(self.upload_button)

        main_layout.addLayout(button_layout)

        # Display Generated Content
        self.result_display = QTextEdit()
        self.result_display.setReadOnly(True)
        self.result_display.setPlaceholderText("Generated content or uploaded file content will appear here...")
        main_layout.addWidget(self.result_display)

    def toggle_custom_tone(self):
        if self.tone_combo.currentText() == "custom":
            self.custom_tone_input.setEnabled(True)
        else:
            self.custom_tone_input.setEnabled(False)

    def toggle_word_limit(self):
        if self.length_combo.currentText() == "Custom":
            self.word_limit_spinbox.setEnabled(True)
        else:
            self.word_limit_spinbox.setEnabled(False)

    def generate_text(self):
        prompt = self.prompt_input.toPlainText().strip()
        tone = self.tone_combo.currentText()
        if tone == "custom":
            tone = self.custom_tone_input.text().strip()

        length = self.length_combo.currentText()
        word_limit = None
        if length == "Custom":
            word_limit = self.word_limit_spinbox.value()

        if not prompt:
            QMessageBox.critical(self, "Error", "Prompt cannot be empty.")
            return

        self.result_display.setText("Generating content...")
        content = generate_content(prompt, tone, length, word_limit)

        if "Error" in content:
            QMessageBox.critical(self, "Error", content)
            self.result_display.setText("")
            return

        self.generated_content = content
        self.result_display.setText(content)
        self.save_button.setEnabled(True)

    def save_text(self):
        if not self.generated_content:
            QMessageBox.critical(self, "Error", "No content to save.")
            return

        save_file, _ = QFileDialog.getSaveFileName(self, "Save File", "", "Text Files (*.txt);;Word Files (*.docx);;PDF Files (*.pdf)")

        if not save_file:
            return

        if save_file.endswith(".txt"):
            save_to_txt(self.generated_content, save_file)
        elif save_file.endswith(".docx"):
            save_to_docx(self.generated_content, save_file)
        elif save_file.endswith(".pdf"):
            save_to_pdf(self.generated_content, save_file)
        else:
            QMessageBox.critical(self, "Error", "Unsupported file type.")
            return

        QMessageBox.information(self, "Success", f"File saved as {save_file}")

    def upload_file(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Open File", "", "Text Files (*.txt);;Word Files (*.docx);;PDF Files (*.pdf)")

        if not file_path:
            return

        file_content = read_file(file_path)

        if "Error" in file_content:
            QMessageBox.critical(self, "Error", file_content)
            return

        self.result_display.setText(file_content)
        self.generated_content = file_content
        self.save_button.setEnabled(True)

# Main Execution
if __name__ == "__main__":
    import sys
    app = QApplication(sys.argv)
    window = AIDocumentAgent()
    window.show()
    sys.exit(app.exec())
