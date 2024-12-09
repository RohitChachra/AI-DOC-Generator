import tkinter as tk
from tkinter import filedialog, messagebox
from huggingface_hub import InferenceClient
import docx
from fpdf import FPDF
from PyPDF2 import PdfReader
from docx import Document

# Hugging Face API Key
API_KEY = ""
client = InferenceClient(api_key=API_KEY)

# Functions
def generate_content(prompt, tone, length, word_limit=None, max_tokens=500):
    """Generate content using Hugging Face API."""
    if length == "Custom" and word_limit:
        formatted_prompt = f"Write in a {tone} tone and limit to {word_limit} words. {prompt}"
    else:
        formatted_prompt = f"Write in a {tone} tone and keep it {length}. {prompt}"
    messages = [{"role": "user", "content": formatted_prompt}]
    try:
        completion = client.chat.completions.create(
            model="meta-llama/Meta-Llama-3-8B-Instruct",
            messages=messages,
            max_tokens=max_tokens,
        )
        return completion.choices[0].message.content
    except Exception as e:
        return f"Error generating content: {e}"

def save_to_txt(content, filename):
    with open(filename, "w", encoding="utf-8") as file:
        file.write(content)
    return filename

def save_to_docx(content, filename):
    doc = docx.Document()
    doc.add_paragraph(content)
    doc.save(filename)
    return filename

def save_to_pdf(content, filename):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, content)
    pdf.output(filename)
    return filename

def convert_docx_to_pdf(input_file, output_file):
    try:
        doc = Document(input_file)
        content = "\n".join([para.text for para in doc.paragraphs])
        save_to_pdf(content, output_file)
        return output_file
    except Exception as e:
        return f"Error converting DOCX to PDF: {e}"

def convert_pdf_to_txt(input_file, output_file):
    try:
        reader = PdfReader(input_file)
        content = "\n".join([page.extract_text() for page in reader.pages])
        save_to_txt(content, output_file)
        return output_file
    except Exception as e:
        return f"Error converting PDF to TXT: {e}"

def convert_pdf_to_docx(input_file, output_file):
    try:
        reader = PdfReader(input_file)
        doc = Document()
        for page in reader.pages:
            doc.add_paragraph(page.extract_text())
        doc.save(output_file)
        return output_file
    except Exception as e:
        return f"Error converting PDF to DOCX: {e}"

# UI Functions
def show_main_screen():
    """Show the main screen."""
    clear_frame()
    tk.Label(root, text="Welcome to the AI Document Agent!", font=("Arial", 16)).pack(pady=10)
    tk.Label(root, text="This agent can generate documents using AI or convert files between formats.", wraplength=400).pack(pady=10)
    tk.Button(root, text="Generate a Document with AI", command=show_generate_screen, width=30, height=2, font=("Arial", 12, "bold")).pack(pady=5)
    tk.Button(root, text="Convert a File", command=show_convert_screen, width=30, height=2, font=("Arial", 12, "bold")).pack(pady=5)

def show_generate_screen():
    """Show the document generation screen."""
    clear_frame()

    def toggle_word_limit(*args):
        """Show or hide the word limit entry box based on selection."""
        if length_var.get() == "Custom":
            word_limit_label.pack(pady=5)
            word_limit_entry.pack(pady=5)
        else:
            word_limit_label.pack_forget()
            word_limit_entry.pack_forget()

    def toggle_custom_tone(*args):
        """Show or hide the custom tone entry box based on selection."""
        if tone_var.get() == "custom":
            custom_tone_label.pack(pady=5)
            custom_tone_entry.pack(pady=5)
        else:
            custom_tone_label.pack_forget()
            custom_tone_entry.pack_forget()

    def generate_and_save():
        prompt = prompt_entry.get("1.0", "end-1c").strip()
        tone = tone_var.get()
        length = length_var.get()
        word_limit = None
        if length == "word limit":
            word_limit = word_limit_entry.get().strip()
            if not word_limit.isdigit():
                messagebox.showerror("Error", "Word limit must be a number.")
                return
        if tone == "custom":
            tone = custom_tone_entry.get().strip()
            if not tone:
                messagebox.showerror("Error", "Custom tone cannot be empty.")
                return
        file_type = file_type_var.get()

        if not prompt:
            messagebox.showerror("Error", "Prompt cannot be empty.")
            return

        # Generate content
        result_label.config(text="Generating content...")
        content = generate_content(prompt, tone, length, word_limit)

        if "Error" in content:
            messagebox.showerror("Error", content)
            result_label.config(text="")
            return

        # Save content
        save_file = filedialog.asksaveasfilename(
            defaultextension=f".{file_type.lower()}",
            filetypes=[(f"{file_type} files", f"*.{file_type.lower()}")],
        )
        if not save_file:
            result_label.config(text="Operation canceled.")
            return

        if file_type == "TXT":
            save_to_txt(content, save_file)
        elif file_type == "DOCX":
            save_to_docx(content, save_file)
        elif file_type == "PDF":
            save_to_pdf(content, save_file)

        messagebox.showinfo("Success", f"File saved as {save_file}")
        result_label.config(text="Content generated and saved.")

    tk.Label(root, text="Generate a Document with AI", font=("Arial", 16)).pack(pady=10)
    tk.Label(root, text="Enter your prompt:").pack()
    prompt_entry = tk.Text(root, height=5, wrap="word")
    prompt_entry.pack(pady=5)

    tk.Label(root, text="Select file type:").pack()
    file_type_var = tk.StringVar(value="TXT")
    tk.OptionMenu(root, file_type_var, "TXT", "DOCX", "PDF").pack()
    
    tk.Label(root, text="Select tone:").pack()
    tone_var = tk.StringVar(value="formal")
    tone_var.trace_add("write", toggle_custom_tone)
    tk.OptionMenu(root, tone_var, "formal", "casual", "professional", "custom").pack()
    custom_tone_label = tk.Label(root, text="Enter custom tone:")
    custom_tone_entry = tk.Entry(root)

    tk.Label(root, text="Select length:").pack()
    length_var = tk.StringVar(value="medium")
    length_var.trace_add("write", toggle_word_limit)
    tk.OptionMenu(root, length_var, "short", "medium", "long", "Custom").pack()
    word_limit_label = tk.Label(root, text="Enter word limit:")
    word_limit_entry = tk.Entry(root)
    
    tk.Button(root, text="Generate and Save", command=generate_and_save, width=20, height=2, font=("Arial", 12, "bold")).pack(pady=5)
    result_label = tk.Label(root, text="", fg="blue", font=("Arial", 12, "bold"))
    result_label.pack(pady=5)

    tk.Button(root, text="Back", command=show_main_screen, width=15, height=2, font=("Arial", 10, "bold")).pack(pady=5)

def show_convert_screen():
    """Show the file conversion screen with a Select File button."""
    clear_frame()

    input_file_var = tk.StringVar(value="No file selected")

    def select_file():
        """Open a file dialog to select a file and update the file path label."""
        input_file = filedialog.askopenfilename(
            title="Select File",
            filetypes=[("All Files", "*.*"), ("DOCX Files", "*.docx"), ("PDF Files", "*.pdf"), ("Text Files", "*.txt")]
        )
        if input_file:
            input_file_var.set(input_file)

    def convert_file():
        """Convert the selected file based on the chosen conversion type."""
        input_file = input_file_var.get()
        if input_file == "No file selected":
            messagebox.showerror("Error", "Please select a file to convert.")
            return

        file_type = convert_type_var.get()

        output_file = filedialog.asksaveasfilename(
            title="Save Converted File",
            defaultextension=f".{'pdf' if 'DOCX to PDF' in file_type else 'docx' if 'PDF to DOCX' in file_type else 'txt'}",
            filetypes=[("All Files", "*.*")]
        )
        if not output_file:
            result_label.config(text="Operation canceled.")
            return

        # Conversion logic
        if file_type == "DOCX to PDF":
            result = convert_docx_to_pdf(input_file, output_file)
        elif file_type == "PDF to TXT":
            result = convert_pdf_to_txt(input_file, output_file)
        elif file_type == "PDF to DOCX":
            result = convert_pdf_to_docx(input_file, output_file)
        else:
            result = "Unsupported conversion type."

        if "Error" in result:
            messagebox.showerror("Error", result)
        else:
            messagebox.showinfo("Success", f"File converted and saved as {output_file}")
            result_label.config(text="File converted successfully.")

    # UI elements for file conversion
    tk.Label(root, text="Convert a File", font=("Arial", 16)).pack(pady=10)

    # File selection button and display
    tk.Button(root, text="Select File", command=select_file, font=("Arial", 10), width=12).pack(pady=5)
    tk.Label(root, textvariable=input_file_var, wraplength=400, fg="blue",font=("Arial", 12, "bold")).pack(pady=5)

    # Conversion type selection
    tk.Label(root, text="Select conversion type:", font=("Arial", 10)).pack()
    convert_type_var = tk.StringVar(value="DOCX to PDF")
    tk.OptionMenu(root, convert_type_var, "DOCX to PDF", "PDF to TXT", "PDF to DOCX").pack(pady=5)

    # Convert button
    tk.Button(root, text="Convert", command=convert_file, width=15, height=2, font=("Arial", 12, "bold")).pack(pady=10)
    result_label = tk.Label(root, text="", fg="blue")
    result_label.pack(pady=5)

    # Back button
    tk.Button(root, text="Back", command=show_main_screen, width=15, height=2, font=("Arial", 10, "bold")).pack(pady=5)

def clear_frame():
    """Clear all widgets in the root frame."""
    for widget in root.winfo_children():
        widget.destroy()

# Main Window
root = tk.Tk()
root.title("AI Document Agent")
root.geometry("500x700")
show_main_screen()
root.mainloop()
